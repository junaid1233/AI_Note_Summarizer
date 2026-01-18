# app.py
import streamlit as st
import pandas as pd
import numpy as np
import os
import re
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import base64
from io import BytesIO
import json
import tempfile
from pathlib import Path

# Import image processing
from PIL import Image, ImageDraw, ImageFont, ImageOps
import textwrap

# File processing
import fitz  # PyMuPDF
from docx import Document
import markdown

# PDF export
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors

# AI/ML imports
try:
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
    import torch

    TRANSFORMERS_AVAILABLE = True
except:
    TRANSFORMERS_AVAILABLE = False

try:
    import google.generativeai as genai

    GOOGLE_AI_AVAILABLE = True
except:
    GOOGLE_AI_AVAILABLE = False

try:
    import openai

    OPENAI_AVAILABLE = True
except:
    OPENAI_AVAILABLE = False

try:
    import anthropic

    ANTHROPIC_AVAILABLE = True
except:
    ANTHROPIC_AVAILABLE = False

# NLP
try:
    import spacy
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords

    NLP_AVAILABLE = True
except:
    NLP_AVAILABLE = False

# Page config
st.set_page_config(
    page_title="AI Study Assistant - Smart Notes & Flashcards",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS with modern design
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

    * {
        font-family: 'Inter', sans-serif;
    }

    .main-header {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 1rem;
        padding: 1rem;
    }

    .sub-header {
        font-size: 1.8rem;
        font-weight: 600;
        color: #2D3748;
        margin-top: 2rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 3px solid #E2E8F0;
    }

    .section-card {
        background: white;
        border-radius: 16px;
        padding: 24px;
        margin: 16px 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
        border: 1px solid #E2E8F0;
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }

    .section-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 30px rgba(0, 0, 0, 0.12);
    }

    .flashcard-container {
        perspective: 1000px;
        width: 100%;
        height: 400px;
        cursor: pointer;
        margin: 20px auto;
    }

    .flashcard {
        width: 100%;
        height: 100%;
        position: relative;
        transform-style: preserve-3d;
        transition: transform 0.8s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        border-radius: 20px;
        box-shadow: 0 15px 35px rgba(0, 0, 0, 0.2);
    }

    .flashcard.flipped {
        transform: rotateY(180deg);
    }

    .flashcard-front, .flashcard-back {
        position: absolute;
        width: 100%;
        height: 100%;
        backface-visibility: hidden;
        border-radius: 20px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        padding: 40px;
        overflow: hidden;
    }

    .flashcard-front {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }

    .flashcard-back {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        transform: rotateY(180deg);
    }

    .flashcard-question {
        font-size: 2rem;
        font-weight: 600;
        text-align: center;
        margin-bottom: 20px;
        line-height: 1.4;
    }

    .flashcard-answer {
        font-size: 1.8rem;
        text-align: center;
        line-height: 1.6;
    }

    .flashcard-subject {
        position: absolute;
        top: 20px;
        left: 20px;
        background: rgba(255, 255, 255, 0.2);
        padding: 6px 16px;
        border-radius: 20px;
        font-size: 0.9rem;
        font-weight: 500;
    }

    .flashcard-hint {
        position: absolute;
        bottom: 20px;
        color: rgba(255, 255, 255, 0.8);
        font-size: 0.9rem;
    }

    .summary-box {
        background: linear-gradient(135deg, #f6f9fc 0%, #ffffff 100%);
        border-left: 5px solid #667eea;
        padding: 24px;
        border-radius: 12px;
        margin: 16px 0;
        box-shadow: 0 2px 10px rgba(0, 0, 0, 0.05);
    }

    .summary-title {
        font-size: 1.4rem;
        font-weight: 600;
        color: #2D3748;
        margin-bottom: 12px;
    }

    .summary-content {
        font-size: 1.1rem;
        line-height: 1.8;
        color: #4A5568;
    }

    .bullet-point {
        display: flex;
        align-items: flex-start;
        margin: 8px 0;
    }

    .bullet-icon {
        color: #667eea;
        font-weight: bold;
        margin-right: 12px;
        font-size: 1.2rem;
    }

    .tag {
        display: inline-flex;
        align-items: center;
        background: linear-gradient(135deg, #E3F2FD 0%, #E8F5E9 100%);
        color: #2E7D32;
        padding: 6px 16px;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 500;
        margin: 4px;
        border: 1px solid rgba(102, 126, 234, 0.2);
    }

    .progress-container {
        background: #E2E8F0;
        border-radius: 10px;
        overflow: hidden;
        margin: 20px 0;
        height: 12px;
    }

    .progress-bar {
        height: 100%;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        border-radius: 10px;
        transition: width 0.5s ease;
    }

    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 14px 28px;
        font-size: 1rem;
        font-weight: 600;
        border-radius: 12px;
        transition: all 0.3s ease;
        width: 100%;
    }

    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 25px rgba(102, 126, 234, 0.4);
    }

    .secondary-button {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%) !important;
    }

    .input-box {
        background: white;
        border: 2px solid #E2E8F0;
        border-radius: 12px;
        padding: 20px;
        margin: 16px 0;
        transition: border-color 0.3s ease;
    }

    .input-box:focus-within {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }

    .stat-card {
        background: white;
        border-radius: 16px;
        padding: 24px;
        text-align: center;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
        border: 1px solid #E2E8F0;
    }

    .stat-number {
        font-size: 2.5rem;
        font-weight: 800;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 10px 0;
    }

    .stat-label {
        font-size: 1rem;
        color: #718096;
        font-weight: 500;
    }

    .toggle-container {
        display: flex;
        background: #F7FAFC;
        border-radius: 12px;
        padding: 4px;
        margin: 16px 0;
    }

    .toggle-option {
        flex: 1;
        text-align: center;
        padding: 12px;
        border-radius: 8px;
        cursor: pointer;
        font-weight: 500;
        transition: all 0.3s ease;
    }

    .toggle-option.active {
        background: white;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        color: #667eea;
    }

    /* Animation for new cards */
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    .slide-in {
        animation: slideIn 0.5s ease;
    }

    /* Responsive design */
    @media (max-width: 768px) {
        .main-header {
            font-size: 2rem;
        }

        .flashcard-question {
            font-size: 1.5rem;
        }

        .flashcard-answer {
            font-size: 1.3rem;
        }
    }
</style>
""", unsafe_allow_html=True)


class AdvancedNoteProcessor:
    def __init__(self):
        self.summarizer = None
        self.nlp = None
        self.init_models()

    def init_models(self):
        """Initialize AI models"""
        # Try to load spaCy model for NLP
        if NLP_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except:
                try:
                    spacy.cli.download("en_core_web_sm")
                    self.nlp = spacy.load("en_core_web_sm")
                except:
                    self.nlp = None

        # Initialize transformers if available
        if TRANSFORMERS_AVAILABLE:
            try:
                self.summarizer = pipeline(
                    "summarization",
                    model="facebook/bart-large-cnn",
                    tokenizer="facebook/bart-large-cnn",
                    device=-1  # Use CPU by default
                )
            except:
                self.summarizer = None

    def extract_text_advanced(self, uploaded_file):
        """Extract text with better formatting"""
        text = ""
        filename = uploaded_file.name.lower()

        try:
            if filename.endswith('.pdf'):
                # Use PyMuPDF for better PDF extraction
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name

                doc = fitz.open(tmp_path)
                for page in doc:
                    text += page.get_text("text") + "\n\n"
                doc.close()
                os.unlink(tmp_path)

            elif filename.endswith('.docx'):
                doc = Document(uploaded_file)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n\n"

                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + " "
                        text += "\n"
                    text += "\n"

            elif filename.endswith('.txt'):
                text = uploaded_file.getvalue().decode('utf-8')

            elif filename.endswith('.md'):
                text = uploaded_file.getvalue().decode('utf-8')

        except Exception as e:
            st.error(f"Error reading file: {str(e)}")

        return self.clean_text(text)

    def clean_text(self, text):
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-—]', ' ', text)
        return text.strip()

    def summarize_with_transformers(self, text, max_length=150, min_length=50):
        """Use transformers for summarization"""
        if not self.summarizer or len(text.split()) < 50:
            return self.summarize_with_nlp(text)

        try:
            # Split text into chunks if too long
            chunks = self.split_text_into_chunks(text, max_chunk_size=1000)
            summaries = []

            for chunk in chunks:
                summary = self.summarizer(
                    chunk,
                    max_length=max_length,
                    min_length=min_length,
                    do_sample=False
                )[0]['summary_text']
                summaries.append(summary)

            return " ".join(summaries)
        except:
            return self.summarize_with_nlp(text)

    def summarize_with_nlp(self, text):
        """Extractive summarization using NLP"""
        if not self.nlp or len(text.split()) < 10:
            return text[:500] + "..." if len(text) > 500 else text

        doc = self.nlp(text)
        sentences = [sent.text for sent in doc.sents]

        if len(sentences) <= 3:
            return text

        # Calculate sentence importance (simple heuristic)
        important_sentences = []
        seen_words = set()

        for sent in sentences:
            words = [token.text.lower() for token in self.nlp(sent)
                     if not token.is_stop and token.is_alpha]
            new_words = [w for w in words if w not in seen_words]

            if new_words:
                important_sentences.append(sent)
                seen_words.update(new_words)

            if len(important_sentences) >= 5:
                break

        return " ".join(important_sentences)

    def split_text_into_chunks(self, text, max_chunk_size=1000):
        """Split text into manageable chunks"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            if current_size + len(word) + 1 <= max_chunk_size:
                current_chunk.append(word)
                current_size += len(word) + 1
            else:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_size = len(word)

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def generate_qa_pairs(self, text, num_pairs=5):
        """Generate question-answer pairs from text"""
        if not self.nlp:
            return []

        doc = self.nlp(text)
        sentences = [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]

        qa_pairs = []
        for i, sentence in enumerate(sentences[:num_pairs * 2]):
            if i % 2 == 0 and i + 1 < len(sentences):
                # Use consecutive sentences as Q&A
                question = self.convert_to_question(sentence)
                answer = sentences[i + 1]
                qa_pairs.append((question, answer))

        return qa_pairs

    def convert_to_question(self, sentence):
        """Convert statement to question"""
        sentence = sentence.strip()
        if sentence.endswith('.'):
            sentence = sentence[:-1]

        # Simple conversion rules
        if sentence.startswith(('The', 'A', 'An', 'This', 'That')):
            sentence = sentence.split(' ', 1)[1]

        return f"What is {sentence.lower()}?" if not sentence.startswith(
            ('What', 'Why', 'How', 'When', 'Where')) else sentence

    def extract_key_concepts(self, text, max_concepts=10):
        """Extract key concepts from text"""
        if not self.nlp:
            return []

        doc = self.nlp(text)
        concepts = []

        for chunk in doc.noun_chunks:
            concept = chunk.text.lower()
            if len(concept.split()) <= 3 and concept not in concepts:
                concepts.append(concept)
                if len(concepts) >= max_concepts:
                    break

        return concepts


class FlashcardGenerator:
    def __init__(self):
        self.styles = {
            "science": {"bg_color": (52, 152, 219), "text_color": (255, 255, 255)},
            "math": {"bg_color": (231, 76, 60), "text_color": (255, 255, 255)},
            "history": {"bg_color": (241, 196, 15), "text_color": (0, 0, 0)},
            "literature": {"bg_color": (155, 89, 182), "text_color": (255, 255, 255)},
            "programming": {"bg_color": (46, 204, 113), "text_color": (255, 255, 255)},
            "general": {"bg_color": (52, 73, 94), "text_color": (255, 255, 255)},
        }

    def create_flashcard_image(self, front_text, back_text, subject="general", style="modern"):
        """Create beautiful flashcard images"""
        width, height = 800, 500

        # Create front card
        front_img = Image.new('RGB', (width, height),
                              color=self.styles.get(subject, self.styles["general"])["bg_color"])
        draw = ImageDraw.Draw(front_img)

        try:
            title_font = ImageFont.truetype("arial.ttf", 32)
            text_font = ImageFont.truetype("arial.ttf", 28)
            small_font = ImageFont.truetype("arial.ttf", 20)
        except:
            title_font = ImageFont.load_default()
            text_font = ImageFont.load_default()
            small_font = ImageFont.load_default()

        # Draw subject
        draw.text((40, 40), subject.upper(),
                  fill=self.styles.get(subject, self.styles["general"])["text_color"],
                  font=small_font)

        # Draw decorative elements
        self.draw_decoration(draw, width, height)

        # Draw main text
        self.draw_wrapped_text(draw, front_text, (width // 2, height // 2 - 20),
                               width - 80, text_font,
                               self.styles.get(subject, self.styles["general"])["text_color"])

        # Draw hint
        draw.text((width - 200, height - 60), "Click to flip →",
                  fill=(255, 255, 255, 180), font=small_font)

        # Create back card
        back_img = Image.new('RGB', (width, height), color=(30, 30, 40))
        draw_back = ImageDraw.Draw(back_img)

        # Draw subject
        draw_back.text((40, 40), f"ANSWER - {subject.upper()}",
                       fill=(255, 255, 255), font=small_font)

        # Draw decorative elements
        self.draw_decoration(draw_back, width, height, is_back=True)

        # Draw answer text
        self.draw_wrapped_text(draw_back, back_text, (width // 2, height // 2),
                               width - 80, text_font, (255, 255, 255))

        # Draw hint
        draw_back.text((60, height - 60), "← Back to question",
                       fill=(255, 255, 255, 180), font=small_font)

        return front_img, back_img

    def draw_decoration(self, draw, width, height, is_back=False):
        """Draw decorative elements"""
        # Draw circles
        for i in range(3):
            x = width // 4 * i + 100
            y = height // 4
            r = 40
            draw.ellipse([x - r, y - r, x + r, y + r],
                         outline=(255, 255, 255, 50) if not is_back else (255, 255, 255, 30),
                         width=2)

    def draw_wrapped_text(self, draw, text, center, max_width, font, color):
        """Draw wrapped text centered"""
        lines = textwrap.wrap(text, width=40)
        total_height = len(lines) * 40
        y = center[1] - total_height // 2

        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            text_width = bbox[2] - bbox[0]
            x = center[0] - text_width // 2
            draw.text((x, y), line, fill=color, font=font)
            y += 40


class StudySessionManager:
    def __init__(self):
        self.session_data = {
            "total_cards_studied": 0,
            "mastered_cards": 0,
            "session_start": datetime.now(),
            "cards_per_subject": {},
            "study_time": 0
        }

    def update_progress(self, card_id, mastered=False):
        """Update study progress"""
        self.session_data["total_cards_studied"] += 1
        if mastered:
            self.session_data["mastered_cards"] += 1

    def get_stats(self):
        """Get study statistics"""
        return {
            "total_studied": self.session_data["total_cards_studied"],
            "mastered": self.session_data["mastered_cards"],
            "accuracy": (self.session_data["mastered_cards"] /
                         max(self.session_data["total_cards_studied"], 1)) * 100,
            "time_elapsed": (datetime.now() - self.session_data["session_start"]).seconds // 60
        }


def main():
    # Initialize session state
    if 'processor' not in st.session_state:
        st.session_state.processor = AdvancedNoteProcessor()

    if 'flashcard_gen' not in st.session_state:
        st.session_state.flashcard_gen = FlashcardGenerator()

    if 'study_manager' not in st.session_state:
        st.session_state.study_manager = StudySessionManager()

    if 'summaries' not in st.session_state:
        st.session_state.summaries = []

    if 'flashcards' not in st.session_state:
        st.session_state.flashcards = []

    if 'current_card' not in st.session_state:
        st.session_state.current_card = 0

    if 'flipped' not in st.session_state:
        st.session_state.flipped = False

    if 'subjects' not in st.session_state:
        st.session_state.subjects = [
            "General", "Mathematics", "Physics", "Chemistry", "Biology",
            "History", "Geography", "Literature", "Computer Science",
            "Economics", "Psychology", "Philosophy", "Languages"
        ]

    # App header
    st.markdown('<h1 class="main-header">🧠 AI Study Assistant</h1>', unsafe_allow_html=True)
    st.markdown(
        '<p style="text-align: center; color: #718096; font-size: 1.2rem;">Transform your notes into intelligent summaries and interactive flashcards</p>',
        unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.markdown("## ⚙️ **Settings & Controls**")

        # AI Settings
        with st.expander("🤖 AI Configuration", expanded=True):
            summarization_method = st.selectbox(
                "Summarization Method:",
                ["Transformers (BART)", "Extractive NLP", "Key Points", "Bullet Points"]
            )

            summarization_level = st.select_slider(
                "Detail Level:",
                options=["Very Concise", "Concise", "Balanced", "Detailed", "Very Detailed"]
            )

            auto_qa = st.checkbox("Auto-generate Q&A pairs", value=True)
            num_qa_pairs = st.slider("Number of Q&A pairs:", 1, 10, 3, disabled=not auto_qa)

        # Study Settings
        with st.expander("📚 Study Settings"):
            study_mode = st.selectbox(
                "Study Mode:",
                ["Learn New", "Review", "Test Yourself", "Spaced Repetition"]
            )

            cards_per_session = st.slider("Cards per session:", 5, 50, 15)

        # File Upload
        with st.expander("📂 Upload Documents"):
            uploaded_files = st.file_uploader(
                "Drag & drop or select files",
                type=['pdf', 'txt', 'docx', 'md', 'jpg', 'png'],
                accept_multiple_files=True,
                help="Supported: PDF, Word, Text, Markdown, Images"
            )

            if uploaded_files:
                st.success(f"📄 {len(uploaded_files)} file(s) uploaded")

        # Quick Stats
        st.markdown("---")
        st.markdown("### 📊 Quick Stats")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Flashcards", len(st.session_state.flashcards))
        with col2:
            if st.session_state.flashcards:
                mastered = sum(1 for card in st.session_state.flashcards if card.get("mastered", False))
                st.metric("Mastered", mastered)

        # Quick Actions
        st.markdown("---")
        st.markdown("### ⚡ Quick Actions")

        if st.button("🔄 New Study Session", use_container_width=True):
            st.session_state.study_manager = StudySessionManager()
            st.rerun()

        if st.button("📥 Export All", use_container_width=True):
            # Export functionality
            pass

        if st.button("🧹 Clear All", use_container_width=True):
            st.session_state.flashcards = []
            st.session_state.summaries = []
            st.rerun()

    # Main content - Tabs
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Notes", "🎴 Flashcards", "📊 Study", "📈 Analytics"])

    with tab1:
        st.markdown('<div class="section-card slide-in">', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">Note Summarization</h2>', unsafe_allow_html=True)

        col1, col2 = st.columns([3, 1])

        with col1:
            # Text input area
            st.markdown("### Enter or Paste Your Notes")
            note_text = st.text_area(
                "📝",
                height=300,
                placeholder="""Paste your lecture notes, textbook content, or research material here...

For example:
# Biology - Cell Structure
The cell is the basic unit of life. All living organisms are composed of cells.
There are two main types of cells: prokaryotic and eukaryotic.
Prokaryotic cells are simpler and lack a nucleus, while eukaryotic cells have a nucleus and membrane-bound organelles.""",
                label_visibility="collapsed"
            )

        with col2:
            st.markdown("### ✨ Quick Templates")
            templates = {
                "Lecture Notes": "# Topic: [Enter Topic]\n## Key Concepts:\n• \n• \n• \n## Summary:",
                "Textbook Chapter": "# Chapter: [Chapter Title]\n## Main Ideas:\n1. \n2. \n3. \n## Important Terms:",
                "Research Paper": "# Title: [Paper Title]\n## Abstract:\n## Methodology:\n## Results:\n## Conclusion:"
            }

            selected_template = st.selectbox("Choose template:", list(templates.keys()))
            if st.button("Apply Template"):
                note_text = templates[selected_template]

        # Process uploaded files
        if uploaded_files:
            st.markdown("### 📄 Uploaded Files Content")
            for uploaded_file in uploaded_files:
                with st.expander(f"📄 {uploaded_file.name}"):
                    if uploaded_file.name.lower().endswith(('.png', '.jpg', '.jpeg')):
                        st.image(uploaded_file, use_column_width=True)
                        st.info("For images, please extract text separately and paste above.")
                    else:
                        extracted_text = st.session_state.processor.extract_text_advanced(uploaded_file)
                        st.text_area("Extracted Text",
                                     extracted_text[:2000] + "..." if len(extracted_text) > 2000 else extracted_text,
                                     height=200, key=f"extracted_{uploaded_file.name}")
                        if st.button(f"Use this text", key=f"use_{uploaded_file.name}"):
                            note_text = extracted_text

        # Process button
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            subject = st.selectbox("Subject:", st.session_state.subjects, key="note_subject")
        with col2:
            if st.button("🚀 Process & Summarize", use_container_width=True):
                if note_text:
                    with st.spinner("🤖 Processing your notes..."):
                        # Generate summary
                        summary = st.session_state.processor.summarize_with_transformers(note_text)

                        # Extract key concepts
                        concepts = st.session_state.processor.extract_key_concepts(note_text)

                        # Generate Q&A pairs if enabled
                        qa_pairs = []
                        if auto_qa:
                            qa_pairs = st.session_state.processor.generate_qa_pairs(summary, num_qa_pairs)

                        # Store summary
                        summary_data = {
                            "text": note_text[:1000] + "..." if len(note_text) > 1000 else note_text,
                            "summary": summary,
                            "concepts": concepts,
                            "subject": subject,
                            "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "qa_pairs": qa_pairs
                        }
                        st.session_state.summaries.append(summary_data)

                        # Create flashcards from Q&A pairs
                        for i, (question, answer) in enumerate(qa_pairs):
                            front_img, back_img = st.session_state.flashcard_gen.create_flashcard_image(
                                question, answer, subject.lower()
                            )

                            # Convert images to bytes
                            front_buf = BytesIO()
                            back_buf = BytesIO()
                            front_img.save(front_buf, format='PNG')
                            back_img.save(back_buf, format='PNG')

                            flashcard_data = {
                                "id": len(st.session_state.flashcards),
                                "question": question,
                                "answer": answer,
                                "subject": subject,
                                "front_image": front_buf.getvalue(),
                                "back_image": back_buf.getvalue(),
                                "created": datetime.now().isoformat(),
                                "mastered": False,
                                "review_count": 0,
                                "last_reviewed": None
                            }
                            st.session_state.flashcards.append(flashcard_data)

                        st.success(f"✅ Created {len(qa_pairs)} flashcards!")
                        st.rerun()
                else:
                    st.warning("Please enter some text or upload files.")

        # Display recent summaries
        if st.session_state.summaries:
            st.markdown("---")
            st.markdown("### 📋 Recent Summaries")

            for i, summary_data in enumerate(reversed(st.session_state.summaries[-3:])):
                with st.expander(f"📄 Summary {len(st.session_state.summaries) - i} - {summary_data['subject']}",
                                 expanded=i == 0):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"**Date:** {summary_data['date']}")
                        st.markdown(f"**Original length:** {len(summary_data['text'].split())} words")
                        st.markdown(f"**Summary length:** {len(summary_data['summary'].split())} words")
                    with col2:
                        if st.button(f"Create Flashcards", key=f"create_fc_{i}"):
                            # Create flashcards from this summary
                            pass

                    st.markdown('<div class="summary-box">', unsafe_allow_html=True)
                    st.markdown('<div class="summary-title">📌 Summary</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="summary-content">{summary_data["summary"]}</div>', unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

                    # Key concepts
                    if summary_data['concepts']:
                        st.markdown("**🔑 Key Concepts:**")
                        for concept in summary_data['concepts']:
                            st.markdown(f'<span class="tag">{concept}</span>', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

    with tab2:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">Interactive Flashcards</h2>', unsafe_allow_html=True)

        if st.session_state.flashcards:
            # Filter options
            col1, col2, col3 = st.columns(3)
            with col1:
                filter_subject = st.selectbox(
                    "Filter by subject:",
                    ["All"] + list(set(card["subject"] for card in st.session_state.flashcards))
                )
            with col2:
                filter_mastered = st.selectbox(
                    "Show:",
                    ["All cards", "Not mastered", "Mastered only"]
                )
            with col3:
                sort_by = st.selectbox(
                    "Sort by:",
                    ["Newest first", "Oldest first", "Subject", "Review count"]
                )

            # Apply filters
            filtered_cards = st.session_state.flashcards.copy()

            if filter_subject != "All":
                filtered_cards = [card for card in filtered_cards if card["subject"] == filter_subject]

            if filter_mastered == "Not mastered":
                filtered_cards = [card for card in filtered_cards if not card["mastered"]]
            elif filter_mastered == "Mastered only":
                filtered_cards = [card for card in filtered_cards if card["mastered"]]

            # Apply sorting
            if sort_by == "Newest first":
                filtered_cards.sort(key=lambda x: x["created"], reverse=True)
            elif sort_by == "Oldest first":
                filtered_cards.sort(key=lambda x: x["created"])
            elif sort_by == "Subject":
                filtered_cards.sort(key=lambda x: x["subject"])
            elif sort_by == "Review count":
                filtered_cards.sort(key=lambda x: x["review_count"], reverse=True)

            if filtered_cards:
                # Display current flashcard
                current_idx = min(st.session_state.current_card, len(filtered_cards) - 1)
                current_card = filtered_cards[current_idx]

                # Flashcard display
                st.markdown(
                    f'<div class="flashcard-container" onclick="this.querySelector(\'.flashcard\').classList.toggle(\'flipped\')">',
                    unsafe_allow_html=True)
                st.markdown(f'<div class="flashcard {"flipped" if st.session_state.flipped else ""}">',
                            unsafe_allow_html=True)

                # Front
                st.markdown(f'''
                <div class="flashcard-front">
                    <div class="flashcard-subject">{current_card["subject"]}</div>
                    <div class="flashcard-question">{current_card["question"]}</div>
                    <div class="flashcard-hint">Click to reveal answer</div>
                </div>
                ''', unsafe_allow_html=True)

                # Back
                st.markdown(f'''
                <div class="flashcard-back">
                    <div class="flashcard-subject">ANSWER • {current_card["subject"]}</div>
                    <div class="flashcard-answer">{current_card["answer"]}</div>
                    <div class="flashcard-hint">Click to see question again</div>
                </div>
                ''', unsafe_allow_html=True)

                st.markdown('</div></div>', unsafe_allow_html=True)

                # Card controls
                col1, col2, col3, col4 = st.columns([2, 1, 1, 2])
                with col2:
                    if st.button("⬅️ Previous", use_container_width=True, disabled=current_idx == 0):
                        st.session_state.current_card = max(0, current_idx - 1)
                        st.session_state.flipped = False
                        st.rerun()

                with col3:
                    if st.button("Next ➡️", use_container_width=True, disabled=current_idx == len(filtered_cards) - 1):
                        st.session_state.current_card = min(len(filtered_cards) - 1, current_idx + 1)
                        st.session_state.flipped = False
                        st.rerun()

                # Card actions
                st.markdown("---")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    if st.button("✅ Mark as Mastered", use_container_width=True):
                        current_card["mastered"] = True
                        st.session_state.study_manager.update_progress(current_card["id"], mastered=True)
                        st.success("Card marked as mastered!")
                        st.rerun()

                with col2:
                    if st.button("🔄 Need Review", use_container_width=True):
                        current_card["mastered"] = False
                        current_card["review_count"] += 1
                        st.info("Card marked for review")
                        st.rerun()

                with col3:
                    if st.button("✏️ Edit Card", use_container_width=True):
                        # Edit functionality
                        pass

                with col4:
                    if st.button("🗑️ Delete Card", use_container_width=True, type="secondary"):
                        st.session_state.flashcards = [card for card in st.session_state.flashcards if
                                                       card["id"] != current_card["id"]]
                        st.rerun()

                # Progress
                st.markdown("### 📊 Card Progress")
                progress = (current_idx + 1) / len(filtered_cards)
                st.markdown(
                    f'<div class="progress-container"><div class="progress-bar" style="width: {progress * 100}%"></div></div>',
                    unsafe_allow_html=True)
                st.markdown(f"**Card {current_idx + 1} of {len(filtered_cards)}** • {current_card['subject']}")

                # Card stats
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                    st.markdown(f'<div class="stat-number">{current_card["review_count"]}</div>',
                                unsafe_allow_html=True)
                    st.markdown('<div class="stat-label">Reviews</div>', unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

                with col2:
                    status = "✅" if current_card["mastered"] else "📚"
                    st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                    st.markdown(f'<div class="stat-number">{status}</div>', unsafe_allow_html=True)
                    st.markdown('<div class="stat-label">Status</div>', unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

                with col3:
                    st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                    st.markdown(f'<div class="stat-number">{len(current_card["question"].split())}</div>',
                                unsafe_allow_html=True)
                    st.markdown('<div class="stat-label">Words (Q)</div>', unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

                with col4:
                    st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                    created_date = datetime.fromisoformat(current_card["created"]).strftime("%b %d")
                    st.markdown(f'<div class="stat-number">{created_date}</div>', unsafe_allow_html=True)
                    st.markdown('<div class="stat-label">Created</div>', unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

            else:
                st.info("No flashcards match your filters. Try changing your filter settings.")

        else:
            st.info("""
            ## 🎴 No Flashcards Yet!

            To get started:
            1. **Go to the Notes tab** and paste your study material
            2. **Click "Process & Summarize"** to generate flashcards automatically
            3. **Or create custom flashcards** using the form below

            You can also:
            - Upload lecture notes or textbook chapters
            - Use one of our templates for different subjects
            - Import existing notes in various formats
            """)

            # Quick flashcard creation
            st.markdown("---")
            st.markdown("### ✨ Create Custom Flashcard")

            col1, col2 = st.columns(2)
            with col1:
                custom_q = st.text_input("Question:", placeholder="e.g., What is the mitochondria?")
                custom_subject = st.selectbox("Subject:", st.session_state.subjects, key="custom_subject")

            with col2:
                custom_a = st.text_area("Answer:", placeholder="e.g., The powerhouse of the cell", height=100)

            if st.button("Create Flashcard", use_container_width=True) and custom_q and custom_a:
                front_img, back_img = st.session_state.flashcard_gen.create_flashcard_image(
                    custom_q, custom_a, custom_subject.lower()
                )

                front_buf = BytesIO()
                back_buf = BytesIO()
                front_img.save(front_buf, format='PNG')
                back_img.save(back_buf, format='PNG')

                flashcard_data = {
                    "id": len(st.session_state.flashcards),
                    "question": custom_q,
                    "answer": custom_a,
                    "subject": custom_subject,
                    "front_image": front_buf.getvalue(),
                    "back_image": back_buf.getvalue(),
                    "created": datetime.now().isoformat(),
                    "mastered": False,
                    "review_count": 0,
                    "last_reviewed": None
                }
                st.session_state.flashcards.append(flashcard_data)
                st.success("Flashcard created successfully!")
                st.rerun()

        st.markdown('</div>', unsafe_allow_html=True)

    with tab3:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">Study Session</h2>', unsafe_allow_html=True)

        if not st.session_state.flashcards:
            st.info("Create some flashcards first to start studying!")
        else:
            # Study session controls
            col1, col2, col3 = st.columns(3)
            with col1:
                session_type = st.selectbox(
                    "Session Type:",
                    ["Mixed Review", "Focus on Weak Areas", "New Cards Only", "Test Mode"]
                )

            with col2:
                session_duration = st.slider("Duration (minutes):", 5, 60, 25)

            with col3:
                if st.button("▶️ Start Session", use_container_width=True):
                    st.session_state.study_session_active = True
                    st.session_state.session_start_time = datetime.now()
                    st.session_state.session_cards = st.session_state.flashcards.copy()
                    st.rerun()

            # Active study session
            if st.session_state.get('study_session_active', False):
                st.markdown("---")

                # Session timer
                elapsed = (datetime.now() - st.session_state.session_start_time).seconds
                remaining = max(0, session_duration * 60 - elapsed)
                minutes, seconds = divmod(remaining, 60)

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.markdown(f"**⏱️ Time remaining:** {minutes:02d}:{seconds:02d}")

                with col2:
                    cards_studied = st.session_state.study_manager.session_data["total_cards_studied"]
                    total_cards = len(st.session_state.session_cards)
                    progress = cards_studied / total_cards if total_cards > 0 else 0
                    st.markdown(f"**📊 Progress:** {cards_studied}/{total_cards} cards")

                with col3:
                    if st.button("⏹️ End Session", use_container_width=True):
                        st.session_state.study_session_active = False
                        st.rerun()

                # Study interface
                if st.session_state.session_cards:
                    current_study_card = st.session_state.session_cards[0]

                    st.markdown("### 🎯 Current Card")
                    col1, col2 = st.columns([2, 1])

                    with col1:
                        st.markdown(f"**Question:** {current_study_card['question']}")
                        if st.button("Show Answer"):
                            st.markdown(f"**Answer:** {current_study_card['answer']}")

                    with col2:
                        st.markdown("**Rate your recall:**")
                        col_a, col_b, col_c = st.columns(3)
                        with col_a:
                            if st.button("😞 Hard", use_container_width=True):
                                current_study_card["review_count"] += 1
                                st.session_state.session_cards.append(current_study_card)
                                st.session_state.session_cards.pop(0)
                                st.session_state.study_manager.update_progress(current_study_card["id"])
                                st.rerun()

                        with col_b:
                            if st.button("😐 Okay", use_container_width=True):
                                st.session_state.session_cards.pop(0)
                                st.session_state.study_manager.update_progress(current_study_card["id"])
                                st.rerun()

                        with col_c:
                            if st.button("😊 Easy", use_container_width=True):
                                current_study_card["mastered"] = True
                                st.session_state.session_cards.pop(0)
                                st.session_state.study_manager.update_progress(current_study_card["id"], mastered=True)
                                st.rerun()
                else:
                    st.success("🎉 Session completed!")
                    stats = st.session_state.study_manager.get_stats()
                    st.markdown(f"""
                    **Session Summary:**
                    - Total cards studied: {stats['total_studied']}
                    - Cards mastered: {stats['mastered']}
                    - Accuracy: {stats['accuracy']:.1f}%
                    - Time spent: {stats['time_elapsed']} minutes
                    """)

            # Study tips
            st.markdown("---")
            st.markdown("### 💡 Study Tips")

            tips = [
                "**Spaced Repetition**: Review cards at increasing intervals for better long-term retention",
                "**Active Recall**: Try to remember the answer before flipping the card",
                "**Interleaving**: Mix different subjects and topics during study sessions",
                "**Elaboration**: Connect new information to what you already know",
                "**Self-Testing**: Regularly test yourself without looking at the answers"
            ]

            for tip in tips:
                st.markdown(f"• {tip}")

        st.markdown('</div>', unsafe_allow_html=True)

    with tab4:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">Analytics & Progress</h2>', unsafe_allow_html=True)

        if not st.session_state.flashcards:
            st.info("No data available yet. Create some flashcards to see your progress!")
        else:
            # Overall statistics
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                total_cards = len(st.session_state.flashcards)
                st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                st.markdown(f'<div class="stat-number">{total_cards}</div>', unsafe_allow_html=True)
                st.markdown('<div class="stat-label">Total Cards</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            with col2:
                mastered = sum(1 for card in st.session_state.flashcards if card["mastered"])
                st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                st.markdown(f'<div class="stat-number">{mastered}</div>', unsafe_allow_html=True)
                st.markdown('<div class="stat-label">Mastered</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            with col3:
                total_reviews = sum(card["review_count"] for card in st.session_state.flashcards)
                st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                st.markdown(f'<div class="stat-number">{total_reviews}</div>', unsafe_allow_html=True)
                st.markdown('<div class="stat-label">Total Reviews</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            with col4:
                subjects = len(set(card["subject"] for card in st.session_state.flashcards))
                st.markdown('<div class="stat-card">', unsafe_allow_html=True)
                st.markdown(f'<div class="stat-number">{subjects}</div>', unsafe_allow_html=True)
                st.markdown('<div class="stat-label">Subjects</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            # Subject breakdown
            st.markdown("### 📚 By Subject")

            subject_data = {}
            for card in st.session_state.flashcards:
                subject = card["subject"]
                if subject not in subject_data:
                    subject_data[subject] = {"total": 0, "mastered": 0}
                subject_data[subject]["total"] += 1
                if card["mastered"]:
                    subject_data[subject]["mastered"] += 1

            for subject, data in subject_data.items():
                mastery_rate = (data["mastered"] / data["total"]) * 100 if data["total"] > 0 else 0
                st.markdown(f"**{subject}**: {data['mastered']}/{data['total']} mastered ({mastery_rate:.1f}%)")
                st.progress(mastery_rate / 100)

            # Recent activity
            st.markdown("### 📈 Recent Activity")

            # Export options
            st.markdown("---")
            st.markdown("### 📤 Export Data")

            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("Export as PDF", use_container_width=True):
                    # PDF export logic
                    pass

            with col2:
                if st.button("Export as CSV", use_container_width=True):
                    # CSV export logic
                    pass

            with col3:
                if st.button("Backup All Data", use_container_width=True):
                    # Backup logic
                    pass

        st.markdown('</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()